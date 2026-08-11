"""
Shared CV .docx engine — same content, same CV dict, as cv_builder.py's PDF
output, for clients who want an editable Word document rather than (or
alongside) the PDF.

Reads the exact same Clients/<Name>/cv_data.py CV dict as cv_builder.py — no
docx-specific fields are ever added to that schema. See generate_cv.py for
the CLI that builds both formats from one data file.

Like cv_builder.py, this stays ATS-safe: single-column body flow, borderless
layout tables used only for visual columns (never affecting reading order),
a real bordered table only for the one genuinely tabular section. Colour is
used the same way — name, section headers, and the one data-table's header
row only.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1B, 0x2A, 0x4A)
TEAL = RGBColor(0x0E, 0x7C, 0x7B)
GOLD = RGBColor(0xC9, 0xA8, 0x4C)
GRAY = RGBColor(0x55, 0x55, 0x55)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GOLD_HEX = "C9A84C"
NAVY_HEX = "1B2A4A"
MGRAY_HEX = "D0D6E0"

FONT = "Calibri"


def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_bottom_border(paragraph, hex_color=GOLD_HEX, sz=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


def _grid_borders(table, hex_color=MGRAY_HEX):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), hex_color)
        borders.append(el)
    tblPr.append(borders)


def _run(paragraph, text, bold=False, italic=False, size=9.5, color=BLACK, font=FONT):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def _setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(0)
    return doc


def section_header(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _run(p, text.upper(), bold=True, size=11, color=NAVY)
    _set_bottom_border(p)


def render_header(doc, client):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, client["name"], bold=True, size=20, color=NAVY)

    contact_bits = [client["location"], client["email"]]
    if client.get("languages_line"):
        contact_bits.append(client["languages_line"])
    if client.get("links"):
        contact_bits.extend(client["links"])
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    _run(p2, "  ·  ".join(contact_bits), size=9, color=GRAY)


def render_summary(doc, text):
    section_header(doc, "Professional Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _run(p, text, size=9.5)


def render_competencies(doc, items):
    section_header(doc, "Core Competencies")
    half = (len(items) + 1) // 2
    left, right = items[:half], items[half:]
    while len(right) < len(left):
        right.append("")

    table = doc.add_table(rows=len(left), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(table)
    for i, (l_val, r_val) in enumerate(zip(left, right)):
        for col, val in ((0, l_val), (1, r_val)):
            cell = table.cell(i, col)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            if val:
                _run(cell.paragraphs[0], f"•  {val}", size=9)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)


def render_experience(doc, entries, heading="Professional Experience"):
    section_header(doc, heading)
    for e in entries:
        header = doc.add_table(rows=1, cols=2)
        _no_borders(header)
        header.columns[0].width = Cm(12.5)
        header.columns[1].width = Cm(4.5)
        title_cell, dates_cell = header.cell(0, 0), header.cell(0, 1)
        title_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        _run(title_cell.paragraphs[0], e["title"], bold=True, size=10)
        dates_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        dates_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(dates_cell.paragraphs[0], e["dates"], italic=True, size=9, color=GRAY)

        org_p = doc.add_paragraph()
        org_p.paragraph_format.space_after = Pt(3)
        _run(org_p, e["org"], italic=True, size=9.5, color=TEAL)

        for b in e["bullets"]:
            bp = doc.add_paragraph()
            bp.paragraph_format.left_indent = Cm(0.4)
            bp.paragraph_format.space_after = Pt(2)
            _run(bp, f"•  {b}", size=9)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)


def render_condensed_experience(doc, entries, heading="Earlier Experience"):
    section_header(doc, heading)
    for e in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, e["title"], bold=True, size=9)
        _run(p, f", {e['org']} ", size=9)
        _run(p, f"({e['dates']})", italic=True, size=9, color=GRAY)
        _run(p, f" — {e['summary']}", size=9)


def render_table_section(doc, heading, headers, rows, col_ratios=None):
    section_header(doc, heading)
    total_width_cm = 17.2
    if col_ratios is None:
        col_ratios = [1 / len(headers)] * len(headers)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _grid_borders(table)

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        _set_cell_shading(cell, NAVY_HEX)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        _run(cell.paragraphs[0], h, bold=True, size=8.5, color=RGBColor(0xFF, 0xFF, 0xFF))
        cell.width = Cm(total_width_cm * col_ratios[i])

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].paragraph_format.space_after = Pt(0)
            _run(cells[i].paragraphs[0], val, size=8.5)
            cells[i].width = Cm(total_width_cm * col_ratios[i])

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)


def render_plain_list_section(doc, heading, items):
    section_header(doc, heading)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.4)
        p.paragraph_format.space_after = Pt(1)
        _run(p, f"•  {item}", size=9)


def render_references(doc, entries):
    """entries: list of {name, title, org, phone (optional), email (optional)}."""
    section_header(doc, "References")
    for e in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _run(p, e["name"], bold=True, size=9)
        _run(p, f" — {e['title']}, {e['org']}", size=9)

        contact_bits = []
        if e.get("phone"):
            contact_bits.append(f"Tel: {e['phone']}")
        if e.get("email"):
            contact_bits.append(f"Email: {e['email']}")
        if contact_bits:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(4)
            _run(p2, "  ·  ".join(contact_bits), size=9, color=GRAY)


def build_cv_docx(data, output_path):
    doc = _setup_document()
    client = data["client"]

    render_header(doc, client)
    render_summary(doc, data["summary"])
    render_competencies(doc, data["core_competencies"])
    render_experience(doc, data["experience"])
    if data.get("earlier_experience"):
        render_condensed_experience(doc, data["earlier_experience"])
    if data.get("additional_experience"):
        ae = data["additional_experience"]
        render_table_section(doc, ae["heading"], ae["headers"], ae["rows"], ae.get("col_ratios"))
    if data.get("institutional_relationships"):
        render_plain_list_section(
            doc, "Institutional Relationships & Recognition", data["institutional_relationships"],
        )
    render_plain_list_section(doc, "Education", data["education"])
    if data.get("certifications"):
        render_plain_list_section(doc, "Certifications", data["certifications"])
    if data.get("professional_courses"):
        render_plain_list_section(doc, "Professional Courses (Selected)", data["professional_courses"])
    if data.get("references"):
        render_references(doc, data["references"])

    doc.core_properties.title = f"{client['name']} — CV"
    doc.save(output_path)
